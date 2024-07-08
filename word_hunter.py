from docx import Document
import re
import nltk
from nltk.corpus import stopwords,wordnet
import operator

doc = Document('East of Eden.docx')

words = []

# # print(type(doc.paragraphs))

# pattern = r'[ ]'

for paragraph in doc.paragraphs:

    texts = paragraph.text
    texts = texts.replace('(',' ').replace(')',' ').replace('\n',' ').replace('!',' ').replace('.',' ').replace('”',' ').replace(',',' ').replace(':',' ').replace('—',' ').replace('“',' ').replace('?',' ')
    
    words.append(re.split(' ',texts))

# print(words)

word_counter = dict()

nltk.download('stopwords')

common_words = list(stopwords.words('english'))
common_words.append('') #the word list has many unnecessary empty strings (""). I have included this in the common words
                        #so that I can ignore them in the future
filtered_words = []

for paragraph in words:
    for word in paragraph:
        word = word.lower()
        if word not in common_words:
            # print(word)
            # filtered_words.append(word)
            # now i have to check the count of the word in the book
            if word not in word_counter:
                word_counter[word] = 1
            else:
                word_counter[word] += 1
            
sorted_tuples = sorted(word_counter.items(),key=operator.itemgetter(1),reverse=True)
sorted_dict = dict(sorted_tuples)

word_dictionary = dict()

for word in sorted_dict:
    # print(word)
    if(sorted_dict[word]>1):
        definitions = wordnet.synsets(word)
        if definitions:
            i = 0
            meaning = ""
            for x in definitions:
                i += 1
                meaning = meaning + f"{i} . {x.definition()}\n"
                # meaning.

            word_dictionary[word] = meaning
            
        else:
            continue

doc2 = Document()
doc2.add_heading(f"Important words for {doc.core_properties.title} and their meanings\n",level=1)

table = doc2.add_table(rows=len(word_dictionary)+1,cols=2,style='Table Grid')

cols = table.rows[0].cells
cols[0].text = 'Word'
cols[1].text = 'Meaning'

i=1
for word in word_dictionary:
    row = table.rows[i]
    row.cells[0].text = word
    row.cells[1].text = word_dictionary[word]
    i+=1

doc2.save(f'{doc.core_properties.title}_meanings.docx')

# print(word_dictionary)   
 
