import docx

# Create a new document
doc = docx.Document()

# Add a heading
doc.add_heading('Employee Details', level=1)

# Create a table with 1 row and 3 columns
table = doc.add_table(rows=1, cols=3)

# Add table headings
row = table.rows[0].cells
row[0].text = 'Id'
row[1].text = 'Name'
row[2].text = 'Age'

# Add data to the table
data = [(1, 'Alice', 30), (2, 'Bob', 25), (3, 'Charlie', 28)]
for id, name, age in data:
    row = table.add_row().cells
    row[0].text = str(id)
    row[1].text = name
    row[2].text = str(age)

# Save the document
doc.save('employee_details.docx')
